"""Word document processor implementation.

This module provides functionality for processing Word documents, including:
- Text extraction with formatting preservation
- Table extraction
- Header/footer handling
- Image metadata tracking
- Error handling
"""

import logging
from pathlib import Path
from typing import Dict, List, Optional

import docx
from docx.document import Document

from src.ml.processing.document.base import BaseDocumentProcessor, ProcessingResult
from src.ml.processing.document.config import DocumentProcessingConfig, WordProcessingConfig


class WordProcessor(BaseDocumentProcessor):
    """Processor for Word documents (.docx, .doc).

    Handles extraction of text content with options for preserving formatting,
    extracting tables, and tracking image metadata.

    Args:
        config: Document processing configuration
        logger: Optional logger instance
    """

    def __init__(
        self,
        config: Optional[DocumentProcessingConfig] = None,
        logger: Optional[logging.Logger] = None,
    ):
        super().__init__(config, logger)
        self._word_config = config.word_config if config else WordProcessingConfig()
        self._doc: Optional[Document] = None

    def _initialize_resources(self) -> None:
        """Initialize processor resources."""
        self._doc = None
        self.logger.info("Word processor initialized")

    def _cleanup_resources(self) -> None:
        """Clean up processor resources."""
        self._doc = None
        self.logger.info("Word processor cleaned up")

    def _process_document(self, file_path: Path) -> ProcessingResult:
        """Process a Word document.

        Args:
            file_path: Path to Word file

        Returns:
            ProcessingResult containing extracted content and metadata

        Raises:
            DocumentProcessingError: If processing fails
            DocumentValidationError: If document validation fails
        """
        try:
            # Load document
            self._doc = docx.Document(file_path)

            # Extract content
            content = {
                "text": self._extract_text(),
                "metadata": self._extract_metadata(),
            }

            # Add optional content based on configuration
            if self._word_config.extract_tables:
                content["tables"] = self._extract_tables()
            if self._word_config.extract_headers:
                content["headers"] = self._extract_headers()
            if self._word_config.extract_images:
                content["images"] = self._extract_image_metadata()

            return ProcessingResult(status="success", content=content)

        except Exception as e:
            self.logger.exception(f"Failed to process Word file {file_path}: {e}")
            return ProcessingResult(status="error", errors=[str(e)])

    def _extract_text(self) -> List[Dict[str, str]]:
        """Extract text content from document.

        Returns:
            List of dictionaries containing text content and optional formatting
        """
        paragraphs = []
        for para in self._doc.paragraphs:
            if not para.text.strip():
                continue

            content = {
                "text": para.text,
                "style": para.style.name if self._word_config.preserve_formatting else None,
            }
            paragraphs.append(content)

        return paragraphs

    def _extract_tables(self) -> List[List[List[str]]]:
        """Extract tables from document.

        Returns:
            List of tables, where each table is a list of rows
        """
        tables = []
        for table in self._doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):  # Skip empty rows
                    table_data.append(row_data)
            if table_data:  # Skip empty tables
                tables.append(table_data)
        return tables

    def _extract_headers(self) -> Dict[str, List[str]]:
        """Extract headers and footers.

        Returns:
            Dictionary containing header and footer text
        """
        headers = []
        footers = []
        for section in self._doc.sections:
            if section.header.text.strip():
                headers.append(section.header.text.strip())
            if section.footer.text.strip():
                footers.append(section.footer.text.strip())
        return {"headers": headers, "footers": footers}

    def _extract_image_metadata(self) -> List[Dict[str, any]]:
        """Extract image metadata.

        Returns:
            List of dictionaries containing image metadata
        """
        images = []
        for rel in self._doc.part.rels.values():
            if "image" in rel.reltype:
                if self._word_config.max_image_size > 0:
                    if len(rel._target) > self._word_config.max_image_size:
                        self.logger.warning("Skipping image: exceeds size limit")
                        continue
                images.append(
                    {
                        "filename": rel._target.partname,
                        "content_type": rel._target.content_type,
                        "size": len(rel._target),
                    }
                )
        return images

    def _extract_metadata(self) -> Dict[str, str]:
        """Extract document metadata.

        Returns:
            Dictionary containing document metadata
        """
        core_props = self._doc.core_properties
        return {
            "author": core_props.author or "",
            "created": str(core_props.created or ""),
            "modified": str(core_props.modified or ""),
            "title": core_props.title or "",
            "subject": core_props.subject or "",
            "paragraph_count": len(self._doc.paragraphs),
            "table_count": len(self._doc.tables),
        }
