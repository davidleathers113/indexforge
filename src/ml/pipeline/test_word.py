"""Tests for Word document processor."""

from pathlib import Path

import docx
import pytest

from .config.settings import PipelineConfig
from .word import WordProcessingConfig, WordProcessor


@pytest.fixture
def sample_word(tmp_path: Path) -> Path:
    """Create a sample Word file for testing.

    Args:
        tmp_path: Pytest fixture providing temporary directory path

    Returns:
        Path to the created Word file
    """
    file_path = tmp_path / "test.docx"

    # Create test document
    doc = docx.Document()

    # Add a heading
    doc.add_heading("Test Document", 0)

    # Add some paragraphs with different styles
    doc.add_paragraph("This is a regular paragraph.")

    doc.add_heading("Section 1", level=1)
    doc.add_paragraph("Content in section 1.")

    doc.add_heading("Section 2", level=2)
    doc.add_paragraph("Content in section 2.")

    # Save the document
    doc.save(file_path)

    return file_path


def test_default_initialization():
    """Test Word processor initialization with default config."""
    processor = WordProcessor()
    assert isinstance(processor.config, PipelineConfig)
    assert isinstance(processor.processing_config, WordProcessingConfig)
    assert not processor.is_initialized
    assert len(processor.processed_files) == 0


def test_custom_config():
    """Test Word processor with custom configuration."""
    config = WordProcessingConfig(
        extract_headers=False, preserve_formatting=True, max_paragraph_length=1000
    )
    processor = WordProcessor(processing_config=config)
    assert not processor.processing_config.extract_headers
    assert processor.processing_config.preserve_formatting
    assert processor.processing_config.max_paragraph_length == 1000


def test_process_word_file(sample_word: Path):
    """Test processing of a valid Word file."""
    processor = WordProcessor()

    with processor:
        result = processor.process(sample_word)

    assert isinstance(result, dict)
    assert "content" in result
    assert "metadata" in result
    assert "structure" in result

    # Check metadata
    assert result["metadata"]["file_path"] == str(sample_word)
    assert result["metadata"]["file_type"] == ".docx"
    assert result["metadata"]["total_paragraphs"] > 0
    assert result["metadata"]["has_headers"] is True

    # Check structure
    assert len(result["structure"]) > 0
    header = next(p for p in result["structure"] if p.get("is_header"))
    assert header["header_level"] in {0, 1, 2}

    # Check processed files tracking
    assert sample_word in processor.processed_files


def test_process_specific_formatting(sample_word: Path):
    """Test processing with specific formatting options."""
    config = WordProcessingConfig(extract_headers=True, preserve_formatting=True)
    processor = WordProcessor(processing_config=config)

    with processor:
        result = processor.process(sample_word)

    # Check formatting preservation
    assert any(p.get("style") is not None for p in result["structure"])

    # Check header extraction
    headers = [p for p in result["structure"] if p.get("is_header")]
    assert len(headers) > 0
    assert all("header_level" in h for h in headers)


def test_invalid_file_path():
    """Test processing with invalid file path."""
    processor = WordProcessor()
    with processor:
        with pytest.raises(ValueError, match="File not found"):
            processor.process(Path("nonexistent.docx"))


def test_invalid_file_type(tmp_path: Path):
    """Test processing with invalid file type."""
    invalid_file = tmp_path / "test.txt"
    invalid_file.touch()

    processor = WordProcessor()
    with processor:
        with pytest.raises(ValueError, match="Unsupported file type"):
            processor.process(invalid_file)


def test_paragraph_length_limit(tmp_path: Path):
    """Test paragraph length limit enforcement."""
    file_path = tmp_path / "test.docx"
    doc = docx.Document()
    doc.add_paragraph("x" * 6000)  # Exceeds default limit
    doc.save(file_path)

    processor = WordProcessor()
    with processor:
        with pytest.raises(ValueError, match="Paragraph exceeds maximum length"):
            processor.process(file_path)


def test_cleanup():
    """Test cleanup of processor resources."""
    processor = WordProcessor()
    processor._processed_files.add(Path("test.docx"))

    processor.cleanup()
    assert len(processor.processed_files) == 0
    assert not processor.is_initialized
