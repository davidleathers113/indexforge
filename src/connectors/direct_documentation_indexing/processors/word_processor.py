"""Word document processor for content extraction.

This module provides functionality for processing Word documents (.docx, .doc),
extracting their content, structure, and metadata. It supports extraction of:
- Full text content
- Headers and document structure
- Tables and their data
- Image metadata

The processor focuses on .docx format using the python-docx library.
"""

import logging
from pathlib import Path
from typing import Any

from docx import Document

from .base_processor import BaseProcessor


logger = logging.getLogger(__name__)


class WordProcessor(BaseProcessor):
    """Processor for Word document content extraction.

    This class handles the processing of Word documents, extracting their
    content and structure. It supports configurable extraction of headers,
    tables, and image metadata.

    Attributes:
        SUPPORTED_EXTENSIONS (set): Set of supported file extensions.
        extract_headers (bool): Whether to extract headers separately.
        extract_tables (bool): Whether to extract tables.
        extract_images (bool): Whether to extract image metadata.

    Example:
        ```python
        processor = WordProcessor({
            "extract_headers": True,
            "extract_tables": True,
            "extract_images": False
        })

        if processor.can_process(file_path):
            result = processor.process(file_path)
            if result["status"] == "success":
                content = result["content"]
                headers = content.get("headers")
                tables = content.get("tables")
        ```
    """

    SUPPORTED_EXTENSIONS = {".docx", ".doc"}

    def __init__(self, config: dict[str, Any] | None = None):
        """Initialize the Word processor with configuration.

        Sets up the processor with optional configuration parameters that
        control what content is extracted from the documents.

        Args:
            config: Optional configuration dictionary containing:
                - extract_headers (bool): Extract headers separately
                - extract_tables (bool): Extract tables
                - extract_images (bool): Extract image metadata

        Example:
            ```python
            processor = WordProcessor({
                "extract_headers": True,
                "extract_tables": True,
                "extract_images": False
            })
            ```
        """
        super().__init__(config)
        self.extract_headers = config.get("extract_headers", True)
        self.extract_tables = config.get("extract_tables", True)
        self.extract_images = config.get("extract_images", False)

    def can_process(self, file_path: Path) -> bool:
        """Check if the file can be processed by this processor.

        Determines whether the file has a supported extension (.docx, .doc).
        Note that only .docx files can be fully processed.

        Args:
            file_path: Path to the file to check.

        Returns:
            bool: True if the file extension is supported, False otherwise.

        Example:
            ```python
            if processor.can_process(Path("document.docx")):
                result = processor.process(Path("document.docx"))
            ```
        """
        return file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS

    def process(self, file_path: Path) -> dict[str, Any]:
        """Process a Word document and extract its content.

        Processes the document, extracting full text content and any additional
        components specified in the configuration (headers, tables, images).

        Args:
            file_path: Path to the Word document to process.

        Returns:
            Dict[str, Any]: Processing result containing:
                - status: "success" or "error"
                - content: Dict with document content and metadata
                - error: Error message if processing failed

        Raises:
            ValueError: If file format is not supported (.doc).
            Exception: For any other processing errors.

        Example:
            ```python
            result = processor.process(Path("document.docx"))
            if result["status"] == "success":
                content = result["content"]
                text = content["full_text"]
                headers = content.get("headers")
                tables = content.get("tables")
            else:
                error = result["error"]
            ```
        """
        try:
            metadata = self._get_file_metadata(file_path)

            # Only .docx is supported by python-docx
            if file_path.suffix.lower() != ".docx":
                return {
                    "status": "error",
                    "error": "Only .docx format is supported",
                    "content": {"metadata": metadata},
                }

            doc = Document(file_path)
            content = {"full_text": self._extract_full_text(doc), "metadata": metadata}

            # Extract additional components if configured
            if self.extract_headers:
                content["headers"] = self._extract_headers(doc)

            if self.extract_tables:
                content["tables"] = self._extract_tables(doc)

            if self.extract_images:
                content["images"] = self._extract_image_data(doc)

            return {"status": "success", "content": content}

        except Exception as e:
            logger.error(f"Error processing Word document {file_path}: {e!s}")
            return {"status": "error", "error": str(e), "content": {"metadata": metadata}}

    def _extract_full_text(self, doc: Document) -> str:
        """Extract all text content from the document.

        Extracts and concatenates text from all paragraphs in the document,
        preserving paragraph breaks.

        Args:
            doc: Word document object from python-docx.

        Returns:
            str: Full text content of the document with paragraphs
            separated by newlines.

        Note:
            This method preserves paragraph breaks but does not preserve
            other formatting like bold, italic, or font styles.
        """
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)

    def _extract_headers(self, doc: Document) -> dict[str, Any]:
        """Extract headers and their hierarchical structure.

        Extracts headers from the document, organizing them by their level
        in the document hierarchy (Heading 1, 2, 3).

        Args:
            doc: Word document object from python-docx.

        Returns:
            Dict[str, Any]: Dictionary containing:
                - level_1: List of top-level headers
                - level_2: List of second-level headers
                - level_3: List of third-level headers

        Note:
            Only processes headers up to level 3. Headers with higher
            levels are ignored.
        """
        headers = {"level_1": [], "level_2": [], "level_3": []}

        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith("Heading"):
                level = paragraph.style.name.split()[-1]
                if f"level_{level}" in headers:
                    headers[f"level_{level}"].append(paragraph.text)

        return headers

    def _extract_tables(self, doc: Document) -> list[dict[str, Any]]:
        """Extract tables and their data from the document.

        Processes all tables in the document, extracting their content
        and basic structure information.

        Args:
            doc: Word document object from python-docx.

        Returns:
            List[Dict[str, Any]]: List of tables, each containing:
                - rows: Number of rows in the table
                - columns: Number of columns in the table
                - data: 2D list of cell contents

        Note:
            Cell formatting and styling information is not preserved.
        """
        tables_data = []

        for table in doc.tables:
            data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                data.append(row_data)

            tables_data.append(
                {"rows": len(data), "columns": len(data[0]) if data else 0, "data": data}
            )

        return tables_data

    def _extract_image_data(self, doc: Document) -> list[dict[str, Any]]:
        """Extract metadata for images in the document.

        Extracts basic information about images embedded in the document,
        including filenames and content types.

        Args:
            doc: Word document object from python-docx.

        Returns:
            List[Dict[str, Any]]: List of image metadata, each containing:
                - filename: Name of the image file
                - content_type: Type/extension of the image

        Note:
            This method only extracts metadata, not the actual image content.
            Image binary data is not included in the output.
        """
        images = []
        rels = doc.part.rels

        for rel in rels.values():
            if "image" in rel.target_ref:
                images.append(
                    {
                        "filename": rel.target_ref.split("/")[-1],
                        "content_type": rel.target_ref.split(".")[-1],
                    }
                )

        return images
