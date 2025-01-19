"""Tests for the WordProcessor class."""

import logging
from pathlib import Path

import docx
import pytest

from src.core.processors.word import WordProcessor


@pytest.fixture
def processor(caplog):
    """Fixture for creating a Word processor instance."""
    caplog.set_level(logging.DEBUG)
    return WordProcessor(extract_headers=True, extract_tables=True)


@pytest.fixture
def basic_docx(tmp_path) -> Path:
    """Fixture for creating a basic Word document."""
    file_path = tmp_path / "test.docx"
    doc = docx.Document()
    doc.add_heading("Test Document", 0)
    doc.add_paragraph("This is a test paragraph.")
    doc.save(file_path)
    return file_path


@pytest.fixture
def complex_docx(tmp_path) -> Path:
    """Fixture for creating a complex Word document with tables and headers."""
    file_path = tmp_path / "complex.docx"
    doc = docx.Document()

    # Add headers
    doc.add_heading("Main Title", 0)
    doc.add_heading("Section 1", 1)
    doc.add_paragraph("Content for section 1")
    doc.add_heading("Section 2", 1)
    doc.add_paragraph("Content for section 2")

    # Add table
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    cells = table.rows[0].cells
    cells[0].text = "Header 1"
    cells[1].text = "Header 2"
    cells = table.rows[1].cells
    cells[0].text = "Data 1"
    cells[1].text = "Data 2"

    doc.save(file_path)
    return file_path


def test_initialization(processor, caplog):
    """Test Word processor initialization."""
    assert isinstance(processor, WordProcessor)
    assert processor.logger is not None
    assert processor.extract_headers is True
    assert processor.extract_tables is True
    assert "Initializing Word processor" in caplog.text


def test_can_process_docx(processor, basic_docx):
    """Test Word file processing capability."""
    assert processor.can_process(basic_docx) is True
    assert not processor.can_process(Path("test.txt"))


def test_process_basic_docx_success(processor, basic_docx, caplog):
    """Test successful processing of a basic Word document."""
    result = processor.process_file(basic_docx)
    assert result.status == "success"
    assert isinstance(result.content, dict)
    assert "full_text" in result.content
    assert "Test Document" in result.content["full_text"]
    assert "Processing Word document" in caplog.text


def test_process_complex_docx_success(processor, complex_docx):
    """Test successful processing of a complex Word document."""
    result = processor.process_file(complex_docx)
    assert result.status == "success"
    content = result.content

    # Check headers
    assert "headers" in content
    assert len(content["headers"]) >= 3
    assert any(h["text"] == "Main Title" for h in content["headers"])

    # Check tables
    assert "tables" in content
    assert len(content["tables"]) == 1
    assert len(content["tables"][0]) == 2  # 2 rows
    assert "Header 1" in content["tables"][0][0]


def test_process_invalid_file(processor):
    """Test processing an invalid file."""
    result = processor.process_file(Path("nonexistent.docx"))
    assert result.status == "error"
    assert "File not found" in result.error


def test_selective_extraction(tmp_path):
    """Test selective content extraction based on configuration."""
    processor = WordProcessor(extract_headers=True, extract_tables=False)

    # Create test document
    file_path = tmp_path / "selective.docx"
    doc = docx.Document()
    doc.add_heading("Test Header", 1)
    table = doc.add_table(rows=1, cols=1)
    table.cells[0].text = "Test Data"
    doc.save(file_path)

    result = processor.process_file(file_path)
    assert result.status == "success"
    assert "headers" in result.content
    assert "tables" not in result.content


def test_empty_document(tmp_path):
    """Test processing an empty document."""
    file_path = tmp_path / "empty.docx"
    doc = docx.Document()
    doc.save(file_path)

    processor = WordProcessor()
    result = processor.process_file(file_path)
    assert result.status == "success"
    assert result.content["full_text"].strip() == ""


def test_nested_tables(tmp_path):
    """Test processing document with nested tables."""
    file_path = tmp_path / "nested_tables.docx"
    doc = docx.Document()

    # Create main table
    table1 = doc.add_table(rows=2, cols=2)
    table1.style = "Table Grid"
    cell = table1.cell(0, 0)
    cell.text = "Outer Table"

    # Create nested table in cell (1,1)
    p = table1.cell(1, 1).add_paragraph()
    p._element.append(doc.add_table(rows=1, cols=1)._element)

    doc.save(file_path)

    processor = WordProcessor(extract_tables=True)
    result = processor.process_file(file_path)
    assert result.status == "success"
    assert len(result.content["tables"]) >= 1


def test_document_with_images(tmp_path):
    """Test processing document with images (metadata only)."""
    processor = WordProcessor(extract_images=True)
    file_path = tmp_path / "with_images.docx"
    doc = docx.Document()
    doc.add_paragraph("Document with image placeholder")
    doc.save(file_path)

    result = processor.process_file(file_path)
    assert result.status == "success"
    assert "images" in result.content
