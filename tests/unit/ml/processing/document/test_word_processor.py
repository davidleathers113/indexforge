"""Unit tests for Word document processor.

Tests the functionality of Word document processing, including text extraction,
table handling, headers/footers, and image metadata tracking.
"""

from pathlib import Path
from unittest.mock import Mock, patch

import docx
import pytest
from docx.document import Document
from docx.shared import Inches

from src.ml.processing.document.config import DocumentProcessingConfig, WordProcessingConfig
from src.ml.processing.document.errors import DocumentProcessingError, DocumentValidationError
from src.ml.processing.document.word import WordProcessor


@pytest.fixture
def processor():
    """Create a test processor instance."""
    return WordProcessor()


@pytest.fixture
def config():
    """Create a test configuration."""
    return DocumentProcessingConfig(
        word_config=WordProcessingConfig(
            extract_tables=True,
            extract_headers=True,
            extract_images=True,
            preserve_formatting=True,
            max_image_size=1024 * 1024,  # 1MB
        )
    )


@pytest.fixture
def sample_doc(tmp_path):
    """Create a sample Word document for testing."""
    doc = Document()

    # Add paragraphs with different styles
    doc.add_paragraph("Test paragraph 1", style="Heading 1")
    doc.add_paragraph("Test paragraph 2", style="Normal")

    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Cell 1"
    table.cell(0, 1).text = "Cell 2"
    table.cell(1, 0).text = "Cell 3"
    table.cell(1, 1).text = "Cell 4"

    # Add header and footer
    section = doc.sections[0]
    section.header.paragraphs[0].text = "Test Header"
    section.footer.paragraphs[0].text = "Test Footer"

    # Save document
    doc_path = tmp_path / "test.docx"
    doc.save(doc_path)
    return doc_path


def test_processor_initialization(processor):
    """Test processor initialization."""
    assert not processor.is_initialized
    processor.initialize()
    assert processor.is_initialized
    assert processor._doc is None
    processor.cleanup()
    assert not processor.is_initialized


def test_process_document(processor, sample_doc):
    """Test processing a Word document."""
    processor.initialize()
    result = processor.process(sample_doc)
    processor.cleanup()

    assert result.status == "success"
    assert "text" in result.content
    assert "metadata" in result.content
    assert len(result.content["text"]) == 2  # Two paragraphs


def test_text_extraction_with_formatting(processor, sample_doc):
    """Test text extraction with formatting preservation."""
    processor.config = DocumentProcessingConfig(
        word_config=WordProcessingConfig(preserve_formatting=True)
    )
    processor.initialize()
    result = processor.process(sample_doc)
    processor.cleanup()

    assert result.status == "success"
    assert result.content["text"][0]["style"] == "Heading 1"
    assert result.content["text"][1]["style"] == "Normal"


def test_table_extraction(processor, sample_doc):
    """Test table extraction."""
    processor.config = DocumentProcessingConfig(
        word_config=WordProcessingConfig(extract_tables=True)
    )
    processor.initialize()
    result = processor.process(sample_doc)
    processor.cleanup()

    assert result.status == "success"
    assert "tables" in result.content
    assert len(result.content["tables"]) == 1
    assert result.content["tables"][0][0][0] == "Cell 1"


def test_header_footer_extraction(processor, sample_doc):
    """Test header and footer extraction."""
    processor.config = DocumentProcessingConfig(
        word_config=WordProcessingConfig(extract_headers=True)
    )
    processor.initialize()
    result = processor.process(sample_doc)
    processor.cleanup()

    assert result.status == "success"
    assert "headers" in result.content
    assert result.content["headers"]["headers"][0] == "Test Header"
    assert result.content["headers"]["footers"][0] == "Test Footer"


def test_metadata_extraction(processor, sample_doc):
    """Test metadata extraction."""
    processor.initialize()
    result = processor.process(sample_doc)
    processor.cleanup()

    assert result.status == "success"
    assert "metadata" in result.content
    assert "paragraph_count" in result.content["metadata"]
    assert "table_count" in result.content["metadata"]
    assert result.content["metadata"]["paragraph_count"] == 2
    assert result.content["metadata"]["table_count"] == 1


def test_error_handling(processor, tmp_path):
    """Test error handling during processing."""
    non_word_file = tmp_path / "test.txt"
    non_word_file.touch()

    processor.initialize()
    result = processor.process(non_word_file)
    processor.cleanup()

    assert result.status == "error"
    assert len(result.errors) > 0


def test_file_access_error(processor, tmp_path):
    """Test handling of file access errors."""
    word_file = tmp_path / "test.docx"
    word_file.touch()
    word_file.chmod(0o000)  # Remove all permissions

    processor.initialize()
    result = processor.process(word_file)
    processor.cleanup()

    assert result.status == "error"
    assert any("Permission denied" in error for error in result.errors)

    # Restore permissions for cleanup
    word_file.chmod(0o666)


def test_corrupted_file(processor, tmp_path):
    """Test handling of corrupted Word files."""
    word_file = tmp_path / "test.docx"
    with open(word_file, "wb") as f:
        f.write(b"Not a valid Word file")

    processor.initialize()
    result = processor.process(word_file)
    processor.cleanup()

    assert result.status == "error"
    assert len(result.errors) > 0


def test_image_size_limit(processor, tmp_path):
    """Test image size limit enforcement."""
    doc = Document()
    doc_path = tmp_path / "test.docx"

    # Create a mock image that exceeds size limit
    with patch("docx.parts.image.Image") as mock_image:
        mock_image._target.return_value = b"x" * (1024 * 1024 * 2)  # 2MB
        doc.add_picture("dummy.png", width=Inches(1.0))

    doc.save(doc_path)

    processor.config = DocumentProcessingConfig(
        word_config=WordProcessingConfig(
            extract_images=True, max_image_size=1024 * 1024  # 1MB limit
        )
    )
    processor.initialize()
    result = processor.process(doc_path)
    processor.cleanup()

    assert result.status == "success"
    assert "images" in result.content
    assert len(result.content["images"]) == 0  # Image should be skipped


def test_empty_document(processor, tmp_path):
    """Test processing an empty document."""
    doc = Document()
    doc_path = tmp_path / "empty.docx"
    doc.save(doc_path)

    processor.initialize()
    result = processor.process(doc_path)
    processor.cleanup()

    assert result.status == "success"
    assert len(result.content["text"]) == 0
    assert result.content["metadata"]["paragraph_count"] == 0
    assert result.content["metadata"]["table_count"] == 0
