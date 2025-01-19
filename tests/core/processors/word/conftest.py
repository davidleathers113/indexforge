"""Shared fixtures for Word processor tests."""

from collections.abc import Generator
import logging
from pathlib import Path
from typing import Any

from _pytest.fixtures import FixtureRequest
from _pytest.logging import LogCaptureFixture
from docx import Document
import pytest

from src.core.processors.word import WordProcessor

from .test_utils import ProcessorConfig, TestContent, TestDataConfig


logger = logging.getLogger(__name__)


@pytest.fixture(autouse=True)
def cleanup_test_files(request: FixtureRequest, tmp_path: Path):
    """Ensure test files are cleaned up after each test.

    Args:
        request: Pytest request object
        tmp_path: Temporary directory path
    """
    yield
    logger.info("Cleaning up test files for: %s", request.node.name)
    for file in tmp_path.glob("*.docx"):
        logger.debug("Removing file: %s", file)
        file.unlink()


@pytest.fixture(params=["default", "full_extraction", "headers_only", "tables_only"])
def processor_config(request: FixtureRequest) -> dict[str, Any]:
    """Provide various processor configurations.

    Args:
        request: Pytest request object containing the parameter

    Returns:
        Dict[str, Any]: Processor configuration
    """
    config_type = request.param
    logger.info("Creating processor config: %s", config_type)
    return ProcessorConfig.get_config(config_type)


@pytest.fixture
def word_processor(request: FixtureRequest, caplog: LogCaptureFixture) -> WordProcessor:
    """Provide a default Word processor instance.

    Args:
        request: Pytest request object
        caplog: Fixture for capturing log messages

    Returns:
        WordProcessor: Default processor instance
    """
    caplog.set_level(logging.DEBUG)
    logger.info("Creating default Word processor for test: %s", request.node.name)
    processor = WordProcessor()
    yield processor
    logger.info("Cleaning up Word processor for test: %s", request.node.name)


@pytest.fixture
def configured_processor(
    request: FixtureRequest, processor_config: dict[str, Any], caplog: LogCaptureFixture
) -> WordProcessor:
    """Provide a configured Word processor instance.

    Args:
        request: Pytest request object
        processor_config: Configuration for the processor
        caplog: Fixture for capturing log messages

    Returns:
        WordProcessor: Configured processor instance
    """
    caplog.set_level(logging.DEBUG)
    logger.info(
        "Creating configured Word processor for test: %s with config: %s",
        request.node.name,
        processor_config,
    )
    processor = WordProcessor(processor_config)
    yield processor
    logger.info("Cleaning up configured Word processor for test: %s", request.node.name)


class DocumentBuilder:
    """Builder pattern for creating test Word documents."""

    def __init__(self, path: Path):
        """Initialize builder with document path.

        Args:
            path: Path where the document will be saved
        """
        self.doc = Document()
        self.path = path
        self._validate_path(path)
        logger.debug("Created DocumentBuilder for path: %s", path)

    def _validate_path(self, path: Path) -> None:
        """Validate the document path.

        Args:
            path: Path to validate

        Raises:
            ValueError: If path has invalid extension or parent directory doesn't exist
        """
        if path.suffix.lower() != ".docx":
            raise ValueError("Document path must have .docx extension")
        if not path.parent.exists():
            raise ValueError("Parent directory does not exist")

    def add_headers(self, headers: dict[int, str]) -> "DocumentBuilder":
        """Add headers to document.

        Args:
            headers: Dictionary mapping header levels to text

        Returns:
            DocumentBuilder: Self for method chaining

        Raises:
            ValueError: If header level is invalid
        """
        logger.debug("Adding headers: %s", headers)
        for level, text in headers.items():
            if not (0 < level <= 9):
                raise ValueError(f"Invalid header level: {level}")
            self.doc.add_heading(text, level=level)
        return self

    def add_paragraphs(self, texts: list[str]) -> "DocumentBuilder":
        """Add paragraphs to document.

        Args:
            texts: List of paragraph texts

        Returns:
            DocumentBuilder: Self for method chaining
        """
        logger.debug("Adding paragraphs: %s", texts)
        for text in texts:
            self.doc.add_paragraph(text)
        return self

    def add_table(self, rows: int, cols: int, data: list[list[str]]) -> "DocumentBuilder":
        """Add table to document.

        Args:
            rows: Number of rows
            cols: Number of columns
            data: Table data as list of lists

        Returns:
            DocumentBuilder: Self for method chaining

        Raises:
            ValueError: If data dimensions don't match rows/cols
        """
        logger.debug("Adding table: %sx%s", rows, cols)
        if len(data) != rows or any(len(row) != cols for row in data):
            raise ValueError("Data dimensions don't match specified rows/cols")

        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"

        for i, row_data in enumerate(data):
            for j, cell_data in enumerate(row_data):
                table.rows[i].cells[j].text = cell_data
        return self

    def build(self) -> Path:
        """Save and return document path.

        Returns:
            Path: Path to saved document
        """
        logger.debug("Saving document to: %s", self.path)
        self.doc.save(self.path)
        return self.path


@pytest.fixture
def doc_builder(tmp_path: Path) -> Generator[DocumentBuilder, None, None]:
    """Provide a DocumentBuilder instance.

    Args:
        tmp_path: Temporary directory path

    Returns:
        Generator[DocumentBuilder, None, None]: Function to create DocumentBuilder
    """

    def _create_builder(name: str) -> DocumentBuilder:
        return DocumentBuilder(tmp_path / name)

    yield _create_builder


@pytest.fixture(params=["basic", "minimal", "complex"])
def test_content(request: FixtureRequest) -> TestContent:
    """Provide test content configurations.

    Args:
        request: Pytest request object containing the parameter

    Returns:
        TestContent: Test content configuration
    """
    content_type = request.param
    logger.info("Creating test content: %s", content_type)
    return TestDataConfig.get_test_config(content_type)


@pytest.fixture
def sample_docx(doc_builder, test_content: TestContent, request: FixtureRequest) -> Path:
    """Provide a sample Word document with various content types.

    Args:
        doc_builder: Document builder fixture
        test_content: Test content configuration
        request: Pytest request object

    Returns:
        Path: Path to test Word document
    """
    logger.info("Creating sample Word document for test: %s", request.node.name)

    return (
        doc_builder("test.docx")
        .add_headers(test_content.headers)
        .add_paragraphs(test_content.paragraphs)
        .add_table(
            len(test_content.table_data),
            len(test_content.table_data[0]),
            test_content.table_data,
        )
        .build()
    )
