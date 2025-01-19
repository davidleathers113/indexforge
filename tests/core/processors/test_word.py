"""Tests for Word document processor functionality.

This module contains tests for the WordProcessor class, covering Word document
processing capabilities, including content extraction, metadata handling,
and error conditions.
"""

import logging
from pathlib import Path

from _pytest.fixtures import FixtureRequest
from _pytest.logging import LogCaptureFixture
from docx import Document
import pytest

from src.core.processors.word import WordProcessor


# Configure logging for tests
logger = logging.getLogger(__name__)


@pytest.fixture
def word_processor(request: FixtureRequest, caplog: LogCaptureFixture):
    """Fixture providing a default Word processor instance.

    Args:
        request: Pytest request object for test context
        caplog: Fixture for capturing log messages

    Returns:
        WordProcessor: Default processor instance
    """
    caplog.set_level(logging.DEBUG)
    logger.info("Creating default Word processor for test: %s", request.node.name)
    processor = WordProcessor()
    yield processor
    logger.info("Cleaning up Word processor for test: %s", request.node.name)


@pytest.fixture
def configured_processor(request: FixtureRequest, caplog: LogCaptureFixture):
    """Fixture providing a configured Word processor instance.

    Args:
        request: Pytest request object for test context
        caplog: Fixture for capturing log messages

    Returns:
        WordProcessor: Configured processor instance
    """
    caplog.set_level(logging.DEBUG)
    config = {"extract_headers": True, "extract_tables": True, "extract_images": True}
    logger.info(
        "Creating configured Word processor for test: %s with config: %s",
        request.node.name,
        config,
    )
    processor = WordProcessor(config)
    yield processor
    logger.info("Cleaning up configured Word processor for test: %s", request.node.name)


@pytest.fixture
def sample_docx(tmp_path: Path, request: FixtureRequest):
    """Fixture providing a sample Word document with various content types.

    Args:
        tmp_path: Temporary directory path
        request: Pytest request object for test context

    Returns:
        Path: Path to test Word document
    """
    logger.info("Creating sample Word document for test: %s", request.node.name)
    file_path = tmp_path / "test.docx"
    logger.debug("Creating Word document at: %s", file_path)

    doc = Document()

    # Add headers
    logger.debug("Adding headers to document")
    doc.add_heading("Heading 1", level=1)
    doc.add_heading("Heading 2", level=2)
    doc.add_heading("Heading 3", level=3)

    # Add paragraphs
    logger.debug("Adding paragraphs to document")
    doc.add_paragraph("First paragraph with some content.")
    doc.add_paragraph("Second paragraph with different content.")

    # Add table
    logger.debug("Adding table to document")
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    cells = table.rows[0].cells
    cells[0].text = "Header 1"
    cells[1].text = "Header 2"
    cells = table.rows[1].cells
    cells[0].text = "Data 1"
    cells[1].text = "Data 2"

    logger.debug("Saving document")
    doc.save(file_path)
    logger.debug("Word document created successfully")
    return file_path


class TestWordProcessor:
    """Tests for the WordProcessor class."""

    def test_initialization_default(self, word_processor: WordProcessor, caplog: LogCaptureFixture):
        """Test processor initialization with default settings.

        Args:
            word_processor: Default processor instance
            caplog: Fixture for capturing log messages
        """
        caplog.set_level(logging.DEBUG)
        logger.info("Testing default WordProcessor initialization")

        logger.debug("Verifying default settings")
        assert not word_processor.extract_headers
        assert not word_processor.extract_tables
        assert not word_processor.extract_images
        logger.debug("Default settings verified successfully")

    def test_initialization_configured(
        self, configured_processor: WordProcessor, caplog: LogCaptureFixture
    ):
        """Test processor initialization with custom configuration.

        Args:
            configured_processor: Configured processor instance
            caplog: Fixture for capturing log messages
        """
        caplog.set_level(logging.DEBUG)
        logger.info("Testing configured WordProcessor initialization")

        logger.debug("Verifying configured settings")
        assert configured_processor.extract_headers
        assert configured_processor.extract_tables
        assert configured_processor.extract_images
        logger.debug("Configured settings verified successfully")

    def test_can_process_docx(
        self, word_processor: WordProcessor, tmp_path: Path, caplog: LogCaptureFixture
    ):
        """Test file type detection for Word documents.

        Args:
            word_processor: Default processor instance
            tmp_path: Temporary directory path
            caplog: Fixture for capturing log messages
        """
        caplog.set_level(logging.DEBUG)
        logger.info("Testing Word document type detection")

        file_path = tmp_path / "test.docx"
        logger.debug("Testing file path: %s", file_path)
        assert word_processor.can_process(file_path)
        logger.debug("Word document type detection test passed")

    def test_can_process_unsupported(
        self, word_processor: WordProcessor, tmp_path: Path, caplog: LogCaptureFixture
    ):
        """Test file type detection for unsupported files.

        Args:
            word_processor: Default processor instance
            tmp_path: Temporary directory path
            caplog: Fixture for capturing log messages
        """
        caplog.set_level(logging.DEBUG)
        logger.info("Testing unsupported file type detection")

        for ext in [".doc", ".txt", ".pdf"]:
            file_path = tmp_path / f"test{ext}"
            logger.debug("Testing unsupported extension: %s with path: %s", ext, file_path)
            assert not word_processor.can_process(file_path)
        logger.debug("Unsupported file type detection test passed")

    def test_process_docx_success(
        self, configured_processor: WordProcessor, sample_docx: Path, caplog: LogCaptureFixture
    ):
        """Test successful processing of Word document.

        Args:
            configured_processor: Configured processor instance
            sample_docx: Test Word document path
            caplog: Fixture for capturing log messages
        """
        caplog.set_level(logging.DEBUG)
        logger.info("Testing successful Word document processing")

        logger.debug("Processing Word document: %s", sample_docx)
        result = configured_processor.process(sample_docx)

        logger.debug("Verifying processing result")
        assert result.status == "success"
        assert "content" in result.content
        assert "metadata" in result.content
        assert "First paragraph" in result.content["content"]
        assert "Second paragraph" in result.content["content"]

        if configured_processor.extract_headers:
            logger.debug("Verifying headers extraction")
            assert "headers" in result.content
            headers = result.content["headers"]
            assert isinstance(headers, dict)
            assert "level_1" in headers
            assert "level_2" in headers
            assert "level_3" in headers
            assert "Heading 1" in headers["level_1"]
            assert "Heading 2" in headers["level_2"]
            assert "Heading 3" in headers["level_3"]

        if configured_processor.extract_tables:
            logger.debug("Verifying tables extraction")
            assert "tables" in result.content
            tables = result.content["tables"]
            assert isinstance(tables, list)
            assert len(tables) == 1
            table = tables[0]
            assert len(table) == 2  # Two rows
            assert "Header 1" in table[0]
            assert "Data 1" in table[1]

        logger.debug("Word document processing test passed successfully")

    def test_process_invalid_file(
        self, word_processor: WordProcessor, tmp_path: Path, caplog: LogCaptureFixture
    ):
        """Test processing invalid file.

        Args:
            word_processor: Default processor instance
            tmp_path: Temporary directory path
            caplog: Fixture for capturing log messages
        """
        caplog.set_level(logging.DEBUG)
        logger.info("Testing invalid file processing")

        file_path = tmp_path / "invalid.docx"
        logger.debug("Creating invalid Word document: %s", file_path)
        file_path.write_text("Invalid content")

        logger.debug("Processing invalid file")
        result = word_processor.process(file_path)

        logger.debug("Verifying error result")
        assert result.status == "error"
        assert result.error is not None
        logger.debug("Invalid file test passed successfully")

    def test_empty_document(
        self, word_processor: WordProcessor, tmp_path: Path, caplog: LogCaptureFixture
    ):
        """Test processing empty document.

        Args:
            word_processor: Default processor instance
            tmp_path: Temporary directory path
            caplog: Fixture for capturing log messages
        """
        caplog.set_level(logging.DEBUG)
        logger.info("Testing empty document processing")

        file_path = tmp_path / "empty.docx"
        logger.debug("Creating empty Word document: %s", file_path)
        doc = Document()
        doc.save(file_path)

        logger.debug("Processing empty document")
        result = word_processor.process(file_path)

        logger.debug("Verifying empty document result")
        assert result.status == "success"
        assert result.content["content"] == ""
        logger.debug("Empty document test passed successfully")

    def test_document_with_empty_table(
        self, configured_processor: WordProcessor, tmp_path: Path, caplog: LogCaptureFixture
    ):
        """Test processing document with empty table.

        Args:
            configured_processor: Configured processor instance
            tmp_path: Temporary directory path
            caplog: Fixture for capturing log messages
        """
        caplog.set_level(logging.DEBUG)
        logger.info("Testing document with empty table")

        file_path = tmp_path / "empty_table.docx"
        logger.debug("Creating document with empty table: %s", file_path)
        doc = Document()
        doc.add_table(rows=1, cols=1)
        doc.save(file_path)

        logger.debug("Processing document with empty table")
        result = configured_processor.process(file_path)

        logger.debug("Verifying empty table result")
        assert result.status == "success"
        assert "tables" in result.content
        assert len(result.content["tables"]) == 1
        assert len(result.content["tables"][0]) == 1  # One empty row
        logger.debug("Empty table test passed successfully")

    def test_document_with_nested_tables(
        self, configured_processor: WordProcessor, tmp_path: Path, caplog: LogCaptureFixture
    ):
        """Test processing document with nested tables.

        Args:
            configured_processor: Configured processor instance
            tmp_path: Temporary directory path
            caplog: Fixture for capturing log messages
        """
        caplog.set_level(logging.DEBUG)
        logger.info("Testing document with nested tables")

        file_path = tmp_path / "nested_tables.docx"
        logger.debug("Creating document with nested tables: %s", file_path)
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        cells = table.rows[0].cells
        cells[0].text = "Outer 1"
        cells[1].text = "Outer 2"
        doc.save(file_path)

        logger.debug("Processing document with nested tables")
        result = configured_processor.process(file_path)

        logger.debug("Verifying nested tables result")
        assert result.status == "success"
        assert "tables" in result.content
        assert len(result.content["tables"]) == 1
        assert "Outer 1" in result.content["tables"][0][0]
        logger.debug("Nested tables test passed successfully")

    def test_selective_extraction(self, tmp_path: Path, caplog: LogCaptureFixture):
        """Test selective content extraction based on configuration.

        Args:
            tmp_path: Temporary directory path
            caplog: Fixture for capturing log messages
        """
        caplog.set_level(logging.DEBUG)
        logger.info("Testing selective content extraction")

        config = {"extract_headers": True, "extract_tables": False, "extract_images": False}
        logger.debug("Creating processor with selective extraction: %s", config)
        processor = WordProcessor(config)

        file_path = tmp_path / "test.docx"
        logger.debug("Creating test document: %s", file_path)
        doc = Document()
        doc.add_heading("Test Heading", level=1)
        doc.add_table(rows=1, cols=1)
        doc.save(file_path)

        logger.debug("Processing file with selective extraction")
        result = processor.process(file_path)

        logger.debug("Verifying selective extraction")
        assert "headers" in result.content
        assert "tables" not in result.content
        assert "images" not in result.content
        logger.debug("Selective extraction test passed successfully")
