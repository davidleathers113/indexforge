"""Performance benchmark tests for document processors.

Tests the performance characteristics of document processors, including:
- Processing time for different document sizes
- Memory usage patterns
- Resource utilization under load
- Concurrent processing efficiency
"""

import time
from pathlib import Path
from typing import Dict, List

import pandas as pd
import psutil
import pytest
from docx import Document
from docx.shared import Inches

from src.ml.processing.document.base import ProcessingResult
from src.ml.processing.document.config import (
    DocumentProcessingConfig,
    ExcelProcessingConfig,
    WordProcessingConfig,
)
from src.ml.processing.document.excel import ExcelProcessor
from src.ml.processing.document.word import WordProcessor


@pytest.fixture
def large_files(tmp_path) -> Dict[str, List[Path]]:
    """Create large test files for performance testing."""
    files = {
        "word": [],
        "excel": [],
    }

    # Create large Word documents
    for i in range(3):
        doc = Document()
        # Add substantial content
        for j in range(100):  # 100 paragraphs
            doc.add_paragraph(f"Content {j} " * 20)  # Long paragraphs
            if j % 10 == 0:  # Add tables periodically
                table = doc.add_table(rows=5, cols=5)
                for row in table.rows:
                    for cell in row.cells:
                        cell.text = f"Cell content {j}"
        doc_path = tmp_path / f"large_doc_{i}.docx"
        doc.save(doc_path)
        files["word"].append(doc_path)

    # Create large Excel files
    for i in range(3):
        # Create large DataFrame
        df = pd.DataFrame({f"Col_{j}": range(1000) for j in range(20)})  # 1000 rows  # 20 columns
        excel_path = tmp_path / f"large_excel_{i}.xlsx"
        with pd.ExcelWriter(excel_path) as writer:
            for sheet in range(5):  # 5 sheets
                df.to_excel(writer, sheet_name=f"Sheet_{sheet}", index=False)
        files["excel"].append(excel_path)

    return files


@pytest.fixture
def processors():
    """Create processor instances with performance-oriented configurations."""
    return {
        "word": WordProcessor(
            config=DocumentProcessingConfig(
                word_config=WordProcessingConfig(
                    extract_tables=True,
                    extract_headers=True,
                    extract_images=True,
                )
            )
        ),
        "excel": ExcelProcessor(
            config=DocumentProcessingConfig(
                excel_config=ExcelProcessingConfig(
                    max_rows=10000,
                    max_cols=100,
                )
            )
        ),
    }


def measure_performance(func):
    """Decorator to measure execution time and memory usage."""

    def wrapper(*args, **kwargs):
        process = psutil.Process()
        start_mem = process.memory_info().rss
        start_time = time.time()

        result = func(*args, **kwargs)

        end_time = time.time()
        end_mem = process.memory_info().rss

        performance_metrics = {
            "execution_time": end_time - start_time,
            "memory_delta": end_mem - start_mem,
            "peak_memory": end_mem,
        }

        return result, performance_metrics

    return wrapper


@measure_performance
def process_files(processor, files: List[Path]) -> List[ProcessingResult]:
    """Process a list of files and return results."""
    results = []
    processor.initialize()
    for file_path in files:
        results.append(processor.process(file_path))
    processor.cleanup()
    return results


def test_processing_time(processors, large_files):
    """Test processing time for different document sizes."""
    metrics: Dict[str, Dict] = {}

    for proc_type, processor in processors.items():
        results, perf = process_files(processor, large_files[proc_type])
        metrics[proc_type] = perf

        # Verify successful processing
        assert all(result.status == "success" for result in results)

        # Performance assertions
        assert perf["execution_time"] < 30.0  # Should process within 30 seconds
        assert perf["memory_delta"] < 500 * 1024 * 1024  # Memory increase < 500MB


def test_concurrent_performance(processors, large_files):
    """Test performance under concurrent processing load."""
    from concurrent.futures import ThreadPoolExecutor

    def process_batch(proc_type):
        processor = processors[proc_type]
        return process_files(processor, large_files[proc_type])

    with ThreadPoolExecutor(max_workers=2) as executor:
        futures = {proc_type: executor.submit(process_batch, proc_type) for proc_type in processors}

        results = {}
        for proc_type, future in futures.items():
            results[proc_type] = future.result()

    # Verify results and performance
    for proc_type, (proc_results, metrics) in results.items():
        assert all(result.status == "success" for result in proc_results)
        assert metrics["execution_time"] < 45.0  # Allow 50% more time for concurrent
        assert metrics["memory_delta"] < 750 * 1024 * 1024  # Allow 50% more memory


def test_memory_cleanup(processors, large_files):
    """Test memory cleanup after processing large files."""
    process = psutil.Process()
    initial_memory = process.memory_info().rss

    for proc_type, processor in processors.items():
        # Process files
        results, metrics = process_files(processor, large_files[proc_type])
        assert all(result.status == "success" for result in results)

        # Force cleanup
        processor.cleanup()
        del results

        # Check memory after cleanup
        current_memory = process.memory_info().rss
        memory_diff = current_memory - initial_memory
        assert memory_diff < 50 * 1024 * 1024  # Cleanup should leave < 50MB residual


def test_resource_utilization(processors, large_files, tmp_path):
    """Test CPU and memory utilization patterns."""
    import threading
    from queue import Queue

    metrics_queue = Queue()
    stop_monitoring = threading.Event()

    def monitor_resources():
        process = psutil.Process()
        while not stop_monitoring.is_set():
            metrics_queue.put(
                {
                    "cpu_percent": process.cpu_percent(),
                    "memory_percent": process.memory_percent(),
                }
            )
            time.sleep(0.1)

    # Start monitoring
    monitor_thread = threading.Thread(target=monitor_resources)
    monitor_thread.start()

    try:
        # Process files
        for proc_type, processor in processors.items():
            results, _ = process_files(processor, large_files[proc_type])
            assert all(result.status == "success" for result in results)
    finally:
        # Stop monitoring
        stop_monitoring.set()
        monitor_thread.join()

    # Analyze metrics
    cpu_usage = []
    memory_usage = []
    while not metrics_queue.empty():
        metrics = metrics_queue.get()
        cpu_usage.append(metrics["cpu_percent"])
        memory_usage.append(metrics["memory_percent"])

    # Verify resource utilization
    assert max(cpu_usage) < 90.0  # CPU usage should not spike too high
    assert max(memory_usage) < 75.0  # Memory usage should stay reasonable
