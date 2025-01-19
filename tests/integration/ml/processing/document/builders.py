"""Builder classes for creating test documents.

This module provides builder patterns for creating test documents in various formats,
making test data creation more maintainable and flexible.
"""

from pathlib import Path
from typing import List, Tuple

import pandas as pd
from docx import Document
from docx.shared import Inches


class DocumentBuilder:
    """Builder pattern for creating test Word documents."""

    def __init__(self, tmp_path: Path):
        """Initialize the builder with a temporary path.

        Args:
            tmp_path: pytest fixture providing temporary directory
        """
        self.tmp_path = tmp_path
        self.content: List[str] = []
        self.tables: List[Tuple[int, int, str]] = []

    def with_paragraph(self, text: str) -> "DocumentBuilder":
        """Add a paragraph to the document.

        Args:
            text: Content of the paragraph

        Returns:
            self for method chaining
        """
        self.content.append(text)
        return self

    def with_table(self, rows: int, cols: int, content: str) -> "DocumentBuilder":
        """Add a table to the document.

        Args:
            rows: Number of rows
            cols: Number of columns
            content: Content for the first cell

        Returns:
            self for method chaining
        """
        self.tables.append((rows, cols, content))
        return self

    def build(self, filename: str) -> Path:
        """Build and save the document.

        Args:
            filename: Name of the output file

        Returns:
            Path to the created document
        """
        doc = Document()
        for text in self.content:
            doc.add_paragraph(text)
        for rows, cols, content in self.tables:
            table = doc.add_table(rows=rows, cols=cols)
            table.cell(0, 0).text = content
        path = self.tmp_path / filename
        doc.save(path)
        return path


class ExcelBuilder:
    """Builder pattern for creating test Excel documents."""

    def __init__(self, tmp_path: Path):
        """Initialize the builder with a temporary path.

        Args:
            tmp_path: pytest fixture providing temporary directory
        """
        self.tmp_path = tmp_path
        self.sheets: List[Tuple[str, pd.DataFrame]] = []

    def with_sheet(self, name: str, data: dict) -> "ExcelBuilder":
        """Add a sheet to the Excel file.

        Args:
            name: Name of the sheet
            data: Dictionary of column data

        Returns:
            self for method chaining
        """
        df = pd.DataFrame(data)
        self.sheets.append((name, df))
        return self

    def build(self, filename: str) -> Path:
        """Build and save the Excel file.

        Args:
            filename: Name of the output file

        Returns:
            Path to the created file
        """
        path = self.tmp_path / filename
        with pd.ExcelWriter(path) as writer:
            for name, df in self.sheets:
                df.to_excel(writer, sheet_name=name, index=False)
        return path
